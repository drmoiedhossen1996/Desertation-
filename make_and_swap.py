# -*- coding: utf-8 -*-
"""Regenerate all 12 figures with ONE unified design system (typography + restrained
palette), then replace the images in-place in the .docx without touching captions,
numbering, placement or any prose."""
import os, numpy as np
import matplotlib; matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch, Polygon, Rectangle, Ellipse, Circle, Arc

# ---------------- UNIFIED DESIGN SYSTEM ----------------
plt.rcParams.update({
    "font.family":"DejaVu Sans","font.size":9,
    "axes.edgecolor":"#5D6D7E","axes.linewidth":0.9,
    "xtick.color":"#212121","ytick.color":"#212121","text.color":"#212121",
    "savefig.dpi":220,
})
INK="#212121"
REFf="#EAECEE"; REFe="#5D6D7E"; REFl="#9AA6B2"      # grey  = conventional / reference
FLOW="#2E6DA4"; FLOWf="#D6E2EF"                      # blue  = 4D flow
PATH="#B03A2E"; PATHf="#F2D7D5"                      # red   = pathological / AS
MYO ="#117A65"; MYOf="#D0ECE7"                       # teal  = myocardial / ventricular
BOXf="#EEF1F4"; BOXe="#5D6D7E"                       # neutral schematic box
GRNf="#D5EAD8"; GRNe="#3B7A57"                       # supported (evidence tiers)
GRYf="#ECECEC"; GRYe="#B0B0B0"                       # unsupported / empty
TS_T=10.5; TS_B=8.6; TS_A=8.0; TS_S=7.3
LW_BOX=1.2; LW_ARR=1.8; LW_DAT=2.0
OUT="/projects/sandbox/figs"; os.makedirs(OUT,exist_ok=True)

def box(ax,x,y,w,h,title="",body="",fc=BOXf,ec=BOXe,ts=TS_B,bs=TS_S,tc=INK):
    ax.add_patch(FancyBboxPatch((x,y),w,h,boxstyle="round,pad=0.02,rounding_size=0.06",fc=fc,ec=ec,lw=LW_BOX))
    if title: ax.text(x+w/2,y+h-0.34,title,ha="center",va="center",fontsize=ts,weight="bold",color=tc)
    if body:  ax.text(x+w/2,y+(h-0.6)/2,body,ha="center",va="center",fontsize=bs,color="#333")
def arrow(ax,x1,y1,x2,y2,color=REFe,lw=LW_ARR,style="-|>",ls="-"):
    ax.add_patch(FancyArrowPatch((x1,y1),(x2,y2),arrowstyle=style,mutation_scale=14,lw=lw,color=color,linestyle=ls,shrinkA=0,shrinkB=0))
def plabel(ax,t): ax.set_title(t,fontsize=TS_T,weight="bold",loc="left",color=INK)

# ---------------- FIG 1 : AS & TTE limitations ----------------
def fig1():
    fig,(a,b)=plt.subplots(1,2,figsize=(11,4.7))
    a.set_xlim(0,10); a.set_ylim(0,10); a.axis("off"); plabel(a,"A  Aortic stenosis: obstruction to ventricular consequences")
    seq=[("Aortic valve stenosis","calcification \u00b7 leaflet restriction",PATHf,PATH),
         ("\u2191 LV pressure overload","obstruction to LV outflow",BOXf,BOXe),
         ("Concentric hypertrophy","adaptive wall thickening",BOXf,BOXe),
         ("Fibrosis / dysfunction","maladaptive; adverse outcomes",BOXf,BOXe)]
    y=7.6
    for i,(t,s,fc,ec) in enumerate(seq):
        box(a,2.0,y,6.0,1.55,t,s,fc=fc,ec=ec,ts=TS_B,bs=TS_S)
        if i<3: arrow(a,5.0,y,5.0,y-0.6,color=REFe)
        y-=2.15
    b.set_xlim(0,10); b.set_ylim(0,10); b.axis("off"); plabel(b,"B  Why echocardiographic assessment can mislead")
    b.add_patch(Ellipse((3.0,7.2),3.2,1.7,fc=FLOWf,ec=FLOW,lw=1.5))
    b.add_patch(Circle((3.0,7.2),0.85,fc="none",ec=PATH,lw=1.5,ls=(0,(4,2))))
    b.annotate("",xy=(4.6,7.2),xytext=(1.4,7.2),arrowprops=dict(arrowstyle="<->",color=FLOW,lw=1.1))
    b.text(3.0,8.35,"true elliptical LVOT",ha="center",fontsize=TS_S,color=FLOW)
    b.text(3.0,5.35,"assumed circle\n(single diameter)",ha="center",fontsize=TS_S,color=PATH)
    b.text(6.35,7.2,"Continuity equation assumes a\ncircular LVOT \u2192 a single diameter\nunderestimates area \u2192 AVA error",ha="left",va="center",fontsize=TS_A)
    arrow(b,1.6,3.7,4.0,4.7,color=PATH,lw=1.9)
    arrow(b,1.6,3.4,3.7,3.4,color=REFe,lw=1.5,ls=(0,(4,2)))
    b.text(1.4,4.0,"eccentric jet",fontsize=TS_S,color=PATH,ha="left")
    b.text(1.4,2.9,"Doppler beam",fontsize=TS_S,color=REFe,ha="left")
    b.text(5.6,3.8,"beam\u2013jet misalignment can\nunderestimate V$_{peak}$ (see Figure 6)",ha="left",va="center",fontsize=TS_A,color=PATH)
    # enlarged threshold box
    b.add_patch(FancyBboxPatch((0.3,0.25),9.4,1.55,boxstyle="round,pad=0.03,rounding_size=0.06",fc="#F4F7FA",ec=BOXe,lw=1.3))
    b.text(5.0,1.4,"Severe high-gradient AS",ha="center",va="center",fontsize=9.2,weight="bold",color=INK)
    b.text(5.0,0.72,"V$_{peak}$ \u2265 4.0 m/s      mean gradient \u2265 40 mmHg      AVA \u2264 1.0 cm\u00b2",
           ha="center",va="center",fontsize=9.0,weight="bold",color=BOXe)
    fig.tight_layout(); fig.savefig(f"{OUT}/fig1_as_tte.png",bbox_inches="tight"); plt.close(fig)

# ---------------- FIG 2 : workflow ----------------
def fig2():
    fig,ax=plt.subplots(figsize=(14,3.4)); ax.set_xlim(0,14); ax.set_ylim(0,4); ax.axis("off")
    steps=[("Acquisition","3-directional velocity\nencoding \u00b7 ECG gating \u00b7 VENC",BOXf,BOXe),
           ("Corrections","background phase \u00b7\nconcomitant gradient \u00b7 aliasing",BOXf,BOXe),
           ("Segmentation","aorta / chambers\nof interest",BOXf,BOXe),
           ("Retrospective\nplane placement","peak-velocity plane chosen\nAFTER acquisition",FLOWf,FLOW),
           ("Quantification","V$_{peak}$ \u00b7 pressure\ngradient \u00b7 flow",BOXf,BOXe),
           ("Advanced\nparameters","TKE \u00b7 VEL \u00b7 WSS",BOXf,BOXe)]
    w=2.05; gap=0.25; x=0.15; y=1.1; h=1.9
    for i,(t,s,fc,ec) in enumerate(steps):
        box(ax,x,y,w,h,fc=fc,ec=ec)
        ax.text(x+w/2,y+h-0.42,t,ha="center",va="center",fontsize=TS_B,weight="bold",color=(FLOW if fc==FLOWf else INK))
        ax.text(x+w/2,y+0.60,s,ha="center",va="center",fontsize=TS_S,color="#333")
        if i<len(steps)-1: arrow(ax,x+w,y+h/2,x+w+gap,y+h/2,color=REFe)
        x+=w+gap
    ax.text(0.15+3*(w+gap)+w/2,0.6,"key advantage in aortic stenosis",ha="center",fontsize=TS_S,color=FLOW,weight="bold")
    fig.tight_layout(); fig.savefig(f"{OUT}/fig2_workflow.png",bbox_inches="tight"); plt.close(fig)

# ---------------- FIG 3 : retrospective plane ----------------
def fig3():
    fig,(a,b)=plt.subplots(1,2,figsize=(11,4.6))
    for ax in (a,b): ax.set_xlim(0,10); ax.set_ylim(0,10); ax.axis("off")
    plabel(a,"A  2D phase-contrast")
    a.add_patch(Polygon([(3.4,0.5),(3.4,9.5),(6.6,9.5),(6.6,0.5)],closed=True,fc=REFf,ec=REFl))
    a.add_patch(Rectangle((3.4,5.0),3.2,0.32,fc=REFe,alpha=0.85))
    a.text(6.85,5.2,"single plane\nfixed before scan",ha="left",va="center",fontsize=TS_S,color=REFe)
    arrow(a,5.0,1.0,7.0,9.0,color=PATH,lw=2.3)
    a.text(7.1,9.0,"eccentric jet\nmay be off-plane",ha="left",va="center",fontsize=TS_S,color=PATH)
    a.text(5.0,0.2,"plane orientation must be chosen in advance",ha="center",fontsize=TS_A,color="#555")
    plabel(b,"B  4D flow (volumetric)")
    def cuboid(ax,x,y,w,h,dd):
        ax.add_patch(Polygon([(x,y),(x+w,y),(x+w,y+h),(x,y+h)],closed=True,fc=FLOWf,ec=BOXe,lw=LW_BOX))
        ax.add_patch(Polygon([(x,y+h),(x+dd,y+h+dd*0.6),(x+w+dd,y+h+dd*0.6),(x+w,y+h)],closed=True,fc="#E4EDF6",ec=BOXe,lw=1.0))
        ax.add_patch(Polygon([(x+w,y),(x+w+dd,y+dd*0.6),(x+w+dd,y+h+dd*0.6),(x+w,y+h)],closed=True,fc="#CFDCEB",ec=BOXe,lw=1.0))
    cuboid(b,2.2,2.6,4.6,4.6,1.3)
    b.text(4.5,7.95,"velocity acquired throughout a 3D volume",ha="center",fontsize=TS_S,color="#555")
    for yy,c,lw in [(3.6,REFl,1.0),(4.6,REFl,1.0),(5.6,FLOW,2.6)]:
        b.add_patch(Rectangle((2.2,yy),4.6,0.16,fc=c,alpha=0.95))
    b.text(7.1,5.6,"peak-velocity plane\nchosen AFTER acquisition",ha="left",va="center",fontsize=TS_S,color=FLOW,weight="bold")
    b.text(4.5,1.9,"retrospective plane placement",ha="center",fontsize=TS_B,weight="bold",color=FLOW)
    b.text(4.5,1.3,"reduces dependence on a predefined plane orientation",ha="center",fontsize=TS_S,color="#555")
    fig.tight_layout(); fig.savefig(f"{OUT}/fig3_retrospective_plane.png",bbox_inches="tight"); plt.close(fig)

# ---------------- FIG X : directional conflict ----------------
def figX():
    fig,ax=plt.subplots(figsize=(9.4,3.9)); ax.set_xlim(-3,3); ax.set_ylim(0,10); ax.axis("off")
    ax.set_title("Directional inconsistency: 4D flow versus Doppler peak velocity",fontsize=TS_T,weight="bold",color=INK)
    ax.annotate("",xy=(2.7,5),xytext=(-2.7,5),arrowprops=dict(arrowstyle="<->",lw=2,color=REFe))
    ax.axvline(0,ymin=0.35,ymax=0.65,color=REFl,lw=1); ax.text(0,4.2,"agreement",ha="center",fontsize=TS_S,color="#555")
    ax.text(-2.7,5.6,"4D flow LOWER than Doppler",ha="left",fontsize=TS_B,color=INK,weight="bold")
    ax.text(2.7,5.6,"4D flow HIGHER than Doppler",ha="right",fontsize=TS_B,color=INK,weight="bold")
    def marker(x,ud,label):
        y=6.7 if ud>0 else 3.3
        ax.plot([x],[5],marker="o",ms=11,color=FLOW,zorder=5)
        ax.annotate(label,xy=(x,5),xytext=(x,y),ha="center",fontsize=TS_A,color=INK,
                    arrowprops=dict(arrowstyle="-",color=FLOW,lw=1))
    marker(-1.6,+1,"H\u00e4lv\u00e4\u2079\n4D lower\n(bias \u22121.1 m/s)")
    marker(1.0,-1,"Grafton-Clarke\u2078\n4D higher\n(+0.5 m/s)")
    marker(2.0,+1,"Adriaans\u00b3\n4D higher\n(V$_{peak}$ +16.4%)")
    ax.text(0,1.4,"the net direction differs between studies \u2014 4D flow is not uniformly higher or more accurate",
            ha="center",fontsize=TS_A,color="#555")
    fig.tight_layout(); fig.savefig(f"{OUT}/figX_vpeak_conflict.png",bbox_inches="tight"); plt.close(fig)

# ---------------- FIG 4 : Archer PG (restrained greys + one blue) ----------------
def fig4():
    fig,ax=plt.subplots(figsize=(6.2,4.8))
    labels=["Invasive","4D flow CMR","Doppler TTE"]; means=[50,54,61]; sds=[34,26,32]
    cols=[REFe,FLOW,REFl]; x=np.arange(3)
    ax.bar(x,means,yerr=sds,capsize=6,width=0.6,color=cols,edgecolor=INK,linewidth=1.0,error_kw=dict(lw=1.2,ecolor=INK))
    for i,m in enumerate(means): ax.text(i,m+sds[i]+2,f"{m}\u00b1{sds[i]}",ha="center",fontsize=TS_A)
    ax.plot([0,1],[95,95],color=INK,lw=1); ax.text(0.5,96,"P = 0.67 (ns)",ha="center",fontsize=TS_S)
    ax.plot([0,2],[104,104],color=INK,lw=1); ax.text(1.0,105,"P = 0.0002",ha="center",fontsize=TS_S)
    ax.set_xticks(x); ax.set_xticklabels(labels); ax.set_ylabel("Peak pressure gradient (mmHg)"); ax.set_ylim(0,120)
    ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
    fig.tight_layout(); fig.savefig(f"{OUT}/fig4_archer_pg.png",bbox_inches="tight"); plt.close(fig)

# ---------------- FIG 5 : two mechanisms (keystone) ----------------
def fig5():
    fig,(a,b)=plt.subplots(1,2,figsize=(11,5.2))
    a.set_xlim(0,10); a.set_ylim(0,10); a.axis("off"); plabel(a,"A  Eccentric jet and Doppler beam misalignment")
    a.add_patch(Polygon([(3.2,0.5),(4.3,4.6),(4.3,9.5),(3.4,9.5),(3.4,4.8),(2.4,0.5)],closed=True,fc=REFf,ec=REFl))
    a.add_patch(Polygon([(6.8,0.5),(5.7,4.6),(5.7,9.5),(6.6,9.5),(6.6,4.8),(7.6,0.5)],closed=True,fc=REFf,ec=REFl))
    a.text(5,0.9,"LV",ha="center",fontsize=TS_S,color="#555"); a.text(5,9.1,"aorta",ha="center",fontsize=TS_S,color="#555")
    a.add_patch(Rectangle((4.3,4.5),1.4,0.5,fc=PATH,ec="none",alpha=0.55)); a.text(5,4.2,"stenotic orifice",ha="center",fontsize=TS_S,color=PATH)
    ox,oy=5,4.9
    arrow(a,ox,oy,7.6,9.2,color=PATH,lw=2.6); a.text(7.7,9.3,"eccentric\njet",fontsize=TS_S,color=PATH,ha="left",va="center")
    arrow(a,2.0,0.6,6.2,9.2,color=REFe,lw=2.0,ls=(0,(4,2))); a.text(1.6,0.5,"Doppler\nbeam",fontsize=TS_S,color=REFe,ha="left",va="top")
    a.add_patch(Arc((ox,oy),3.2,3.2,angle=0,theta1=58,theta2=78,color=INK,lw=1.3)); a.text(6.05,7.0,r"$\theta$",fontsize=13,color=INK)
    a.text(5.0,2.6,r"$V_{measured}=V_{true}\cdot\cos\theta$",ha="center",fontsize=TS_B,bbox=dict(boxstyle="round,pad=0.3",fc="white",ec=BOXe))
    a.text(5.0,1.7,"misalignment  \u2192  V$_{peak}$ underestimated",ha="center",fontsize=TS_A,color=PATH,weight="bold")
    plabel(b,"B  Doppler vs catheter pressure gradient")
    t=np.linspace(0,1,400)
    LV=np.where(t<0.6,175*np.sin(np.pi*np.clip(t,0,0.6)/0.6),0)
    m=(t>0.05)&(t<0.7); base=np.clip((t-0.05)/0.65,0,1)
    Ao=np.where(m,120*np.sin(np.pi*base)**1.1,0); Aor=np.where(m,138*np.sin(np.pi*base)**1.1,0)
    b.plot(t,LV,color=MYO,lw=LW_DAT,label="LV pressure")
    b.plot(t,Ao,color=REFe,lw=LW_DAT,label="Aortic pressure (vena contracta)")
    b.plot(t,Aor,color=REFe,lw=1.4,ls=(0,(4,2)),label="Aortic pressure (recovered, downstream)")
    diff=LV-Ao; k=int(np.argmax(diff))
    b.annotate("",xy=(t[k],LV[k]),xytext=(t[k],Ao[k]),arrowprops=dict(arrowstyle="<->",color=PATH,lw=1.8))
    b.text(t[k]-0.02,(LV[k]+Ao[k])/2,"Doppler\nmax instantaneous",ha="right",va="center",fontsize=TS_S,color=PATH)
    iLV=int(np.argmax(LV)); iAo=int(np.argmax(Ao))
    b.annotate("",xy=(0.86,LV[iLV]),xytext=(0.86,Ao[iAo]),arrowprops=dict(arrowstyle="<->",color=MYO,lw=1.8))
    b.text(0.88,(LV[iLV]+Ao[iAo])/2,"catheter\npeak-to-peak",ha="left",va="center",fontsize=TS_S,color=MYO)
    b.plot([t[iLV],0.86],[LV[iLV],LV[iLV]],color=MYO,lw=0.7,ls=":"); b.plot([t[iAo],0.86],[Ao[iAo],Ao[iAo]],color=REFe,lw=0.7,ls=":")
    b.annotate("pressure\nrecovery",xy=(0.42,Aor[int(0.42*399)]),xytext=(0.52,158),fontsize=TS_S,color=REFe,ha="left",
               arrowprops=dict(arrowstyle="-|>",color=REFe,lw=1.1))
    b.set_xlabel("time (systole)"); b.set_ylabel("pressure (mmHg)"); b.set_xlim(0,1.0); b.set_ylim(0,190); b.set_xticks([])
    b.spines["top"].set_visible(False); b.spines["right"].set_visible(False); b.legend(fontsize=6.8,loc="upper right",frameon=False)
    fig.text(0.5,0.015,"Two distinct mechanisms of measurement discrepancy \u2014 they must not be conflated",
             ha="center",fontsize=TS_T,weight="bold",color=INK)
    fig.tight_layout(rect=[0,0.05,1,1]); fig.savefig(f"{OUT}/fig5_two_mechanisms.png",bbox_inches="tight"); plt.close(fig)

# ---------------- FIG 6 : Aliabadi VEL ----------------
def fig6():
    fig,ax=plt.subplots(figsize=(6.6,4.8))
    labels=["TAV\ncontrol","BAV\nno complications","BAV\nregurgitation","BAV\nstenosis"]
    med=[4.1,6.4,11.4,16.2]; lo=[3.4,5.1,9.5,9.1]; hi=[5.7,8.0,17.6,24.4]
    yerr=[[m-l for m,l in zip(med,lo)],[h-m for h,m in zip(hi,med)]]; x=np.arange(4)
    ax.errorbar(x,med,yerr=yerr,fmt="o",ms=9,color=INK,ecolor=REFl,elinewidth=1.6,capsize=6,capthick=1.6)
    for i,mm in enumerate(med): ax.text(i+0.12,mm,f"{mm}",fontsize=TS_A,va="center")
    ax.text(1,8.0+0.7,"ns vs control",ha="center",fontsize=TS_S,color=PATH)
    ax.set_xticks(x); ax.set_xticklabels(labels,fontsize=TS_A); ax.set_ylabel("Peak-systolic viscous energy loss (mW)"); ax.set_ylim(0,27)
    ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
    ax.text(0.02,0.98,"median \u00b7 whiskers = IQR",transform=ax.transAxes,fontsize=TS_S,va="top",color="#555")
    fig.tight_layout(); fig.savefig(f"{OUT}/fig6_aliabadi_vel.png",bbox_inches="tight"); plt.close(fig)

# ---------------- FIG Y : Elhawaz KE ----------------
def figY():
    fig,ax=plt.subplots(figsize=(7.0,4.8))
    comps=["Direct","Delayed","Retained","Residual"]; pre=[4.9,2.46,1.07,0.84]; post=[1.86,1.38,0.91,0.98]; sig=["*","*","ns","ns"]; x=np.arange(4)
    for i in range(4): ax.plot([x[i]-0.13,x[i]+0.13],[pre[i],post[i]],color=REFl,lw=1.5,zorder=1)
    ax.plot(x-0.13,pre,"o",ms=10,color=INK,label="pre-intervention",zorder=3)
    ax.plot(x+0.13,post,"o",ms=10,mfc="white",mec=FLOW,mew=1.9,label="post-intervention",zorder=3)
    for i in range(4): ax.text(x[i],max(pre[i],post[i])+0.35,sig[i],ha="center",fontsize=10,color=(PATH if sig[i]=="*" else "#888"),weight="bold")
    ax.set_xticks(x); ax.set_xticklabels([c+"\nflow" for c in comps]); ax.set_ylabel("LV blood-flow kinetic energy (\u03bcJ)"); ax.set_ylim(0,6.2)
    ax.legend(fontsize=TS_A,frameon=False,loc="upper right"); ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
    ax.text(0.02,0.02,"median values; no error bars (skewed data, reported as median [IQR])",transform=ax.transAxes,fontsize=TS_S,color="#555")
    fig.tight_layout(); fig.savefig(f"{OUT}/figY_elhawaz_ke.png",bbox_inches="tight"); plt.close(fig)

# ---------------- FIG 7 : PV loop (polished) ----------------
def fig7():
    fig,ax=plt.subplots(figsize=(6.4,5.4))
    ESV,EDV,Pao,Pfill,Pes=45,125,120,8,125
    exx=np.linspace(EDV,ESV,80); eyy=np.linspace(Pao,Pes,80)+12*np.sin(np.linspace(0,np.pi,80))
    fxx=np.linspace(ESV,EDV,80); fyy=Pfill+(np.linspace(0,1,80)**2)*(Pfill+2)+2.2*np.sin(np.linspace(0,np.pi,80))
    loopx=np.concatenate([[EDV,EDV],exx,[ESV],fxx]); loopy=np.concatenate([[Pfill,Pao],eyy,[12],fyy])
    ax.fill(loopx,loopy,color=FLOWf,alpha=0.95,zorder=0)
    ax.plot([EDV,EDV],[Pfill,Pao],color=INK,lw=2.1); ax.plot(exx,eyy,color=INK,lw=2.1)
    ax.plot([ESV,ESV],[Pes,12],color=INK,lw=2.1); ax.plot(fxx,fyy,color=INK,lw=2.1)
    ax.plot([20,ESV+3],[0,Pes+8],color=PATH,lw=1.4,ls=(0,(5,3))); ax.text(ESV-2,Pes+14,"ESPVR",color=PATH,fontsize=TS_S)
    exx2=np.linspace(20,140,60); ax.plot(exx2,2+0.0009*(exx2-20)**2.05,color=MYO,lw=1.4,ls=(0,(5,3))); ax.text(131,20,"EDPVR",color=MYO,fontsize=TS_S)
    ax.text((ESV+EDV)/2,70,"stroke\nwork",ha="center",va="center",fontsize=TS_B,weight="bold",color=INK)
    ax.annotate("ejection",xy=(85,eyy[np.argmin(abs(exx-85))]),xytext=(66,152),fontsize=TS_S,arrowprops=dict(arrowstyle="-|>",lw=1,color="#666"))
    ax.text(EDV+2,(Pfill+Pao)/2,"isovolumetric\ncontraction",fontsize=TS_S,color="#666",va="center")
    ax.text(ESV-24,(Pes+12)/2,"isovolumetric\nrelaxation",fontsize=TS_S,color="#666",va="center")
    ax.text((ESV+EDV)/2,4,"filling",ha="center",fontsize=TS_S,color="#666")
    ax.text(78,182,"PVA = stroke work + potential energy",ha="center",fontsize=TS_A,color=PATH,bbox=dict(boxstyle="round,pad=0.3",fc="white",ec=PATH))
    ax.set_xlabel("LV volume (mL)"); ax.set_ylabel("LV pressure (mmHg)"); ax.set_xlim(15,150); ax.set_ylim(0,195)
    ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
    fig.tight_layout(); fig.savefig(f"{OUT}/fig7_pvloop.png",bbox_inches="tight"); plt.close(fig)

# ---------------- FIG 8 : flow energy vs myocardial work ----------------
def fig8():
    fig,ax=plt.subplots(figsize=(11,4.8)); ax.set_xlim(0,12); ax.set_ylim(0,7.5); ax.axis("off")
    box(ax,0.3,1.6,3.2,4.2,fc=FLOWf,ec=FLOW)
    ax.text(1.9,5.35,"Blood-flow energy\n(4D flow CMR)",ha="center",va="center",fontsize=TS_B,weight="bold",color=FLOW)
    ax.text(1.9,3.4,"\u2022 Intracavitary kinetic energy\n\u2022 Turbulent kinetic energy (TKE)\n\u2022 Viscous energy loss (VEL)",ha="center",va="center",fontsize=TS_S)
    box(ax,4.4,1.6,3.2,4.2,fc=PATHf,ec=PATH)
    ax.text(6.0,5.35,"Aortic stenosis",ha="center",va="center",fontsize=TS_B,weight="bold",color=PATH)
    ax.text(6.0,3.4,"\u2022 Valvular obstruction\n\u2022 Altered / eccentric flow\n\u2022 \u2191 LV pressure load",ha="center",va="center",fontsize=TS_S)
    box(ax,8.5,1.6,3.2,4.2,fc=MYOf,ec=MYO)
    ax.text(10.1,5.35,"Myocardial\nmechanical energy",ha="center",va="center",fontsize=TS_B,weight="bold",color=MYO)
    ax.text(10.1,3.35,"\u2022 PV-loop stroke work\n\u2022 Pressure\u2013volume area (PVA)\n\u2022 Ventricular efficiency",ha="center",va="center",fontsize=TS_S)
    arrow(ax,4.4,3.7,3.5,3.7,color=PATH); arrow(ax,7.6,3.7,8.5,3.7,color=PATH)
    ax.text(4.0,4.05,"altered flow",fontsize=TS_S,color=PATH,ha="center"); ax.text(8.05,4.05,"pressure load",fontsize=TS_S,color=PATH,ha="center")
    ax.add_patch(FancyArrowPatch((1.9,1.5),(10.1,1.5),arrowstyle="<->",mutation_scale=16,lw=1.6,color=REFe,connectionstyle="arc3,rad=-0.25"))
    ax.text(6.0,0.35,"physiologically related  \u2014  but NOT interchangeable",ha="center",fontsize=TS_T,weight="bold",color=PATH)
    fig.tight_layout(); fig.savefig(f"{OUT}/fig8_flow_vs_work.png",bbox_inches="tight"); plt.close(fig)

# ---------------- FIG 9 : evidence hierarchy ----------------
def fig9():
    fig,ax=plt.subplots(figsize=(9.6,5.4)); ax.set_xlim(0,12); ax.set_ylim(0,10.5); ax.axis("off")
    tiers=[("1  Measurement validity / agreement","Adriaans\u00b3 \u00b7 Grafton-Clarke\u2076 \u00b7 Archer\u00b9\u00b9",GRNf,GRNe),
           ("2  Association with symptoms / function","Archer\u00b9\u00b9 \u00b7 Elhawaz\u00b9\u2077 (6MWT)",GRNf,GRNe),
           ("3  Association with a clinical decision","Grafton-Clarke\u2078 (intervention; exploratory)",FLOWf,FLOW),
           ("4  Demonstrated clinical utility","no evidence yet",GRYf,GRYe),
           ("5  Improved patient outcomes","no evidence yet",GRYf,GRYe)]
    yb=0.7; h=1.65; cx=6.0
    for i,(t,s,fc,ec) in enumerate(tiers):
        halfw=4.7-0.55*i; y=yb+i*(h+0.12)
        ax.add_patch(FancyBboxPatch((cx-halfw,y),2*halfw,h,boxstyle="round,pad=0.02,rounding_size=0.05",fc=fc,ec=ec,lw=1.4))
        ax.text(cx,y+h*0.62,t,ha="center",va="center",fontsize=TS_B,weight="bold",color=INK)
        col=PATH if s=="no evidence yet" else "#333"
        ax.text(cx,y+h*0.24,s,ha="center",va="center",fontsize=TS_S,color=col,style=("italic" if s=="no evidence yet" else "normal"))
    arrow(ax,0.7,yb,0.7,yb+5*(h+0.12)-0.2,color=REFe)
    ax.text(0.35,yb+2.5*(h+0.12),"increasing clinical maturity",rotation=90,va="center",ha="center",fontsize=TS_A,color=REFe)
    ax.text(cx,yb-0.5,"strongest current evidence at the base; top tiers not yet supported",ha="center",fontsize=TS_A,color="#555")
    fig.tight_layout(); fig.savefig(f"{OUT}/fig9_evidence_hierarchy.png",bbox_inches="tight"); plt.close(fig)

# ---------------- FIG Z : Grafton-Clarke Vpeak by intervention ----------------
def figZ():
    fig,ax=plt.subplots(figsize=(6.8,4.8))
    groups=["Doppler (CWD)","4D flow CMR"]; noiv=[2.6,2.7]; iv=[3.4,4.2]; x=np.arange(2); w=0.34
    ax.bar(x-w/2,noiv,w,label="no intervention",color=REFl,edgecolor=INK,lw=1.0)
    ax.bar(x+w/2,iv,w,label="intervention",color=FLOW,edgecolor=INK,lw=1.0)
    for i in range(2):
        ax.text(x[i]-w/2,noiv[i]+0.05,f"{noiv[i]}",ha="center",fontsize=TS_A)
        ax.text(x[i]+w/2,iv[i]+0.05,f"{iv[i]}",ha="center",fontsize=TS_A)
    ax.text(0,4.5,"P = 0.0025",ha="center",fontsize=TS_S); ax.text(1,4.7,"P < 0.0001",ha="center",fontsize=TS_S)
    ax.axhline(3.5,color=PATH,ls=(0,(5,3)),lw=1.4); ax.text(1.48,3.58,"3.5 m/s (exploratory)",ha="right",fontsize=TS_S,color=PATH)
    ax.set_xticks(x); ax.set_xticklabels(groups); ax.set_ylabel("Peak velocity (m/s)"); ax.set_ylim(0,5.2)
    ax.legend(fontsize=TS_A,frameon=False,loc="upper left"); ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
    fig.tight_layout(); fig.savefig(f"{OUT}/figZ_gc_vpeak.png",bbox_inches="tight"); plt.close(fig)

for f in (fig1,fig2,fig3,figX,fig4,fig5,fig6,figY,fig7,fig8,fig9,figZ):
    f()
print("regenerated all 12 figures")

# ---------------- SWAP images in-place in the .docx ----------------
from docx import Document
from docx.shared import Inches
FN="Final_4D_Flow_CMR_Narrative_Review_FINAL_corrected.docx"
d=Document(FN)
# final figure number -> (file, width in inches)
NUM={1:("fig1_as_tte.png",6.5),2:("fig2_workflow.png",6.5),3:("fig3_retrospective_plane.png",6.5),
     4:("figX_vpeak_conflict.png",6.0),5:("fig4_archer_pg.png",4.7),6:("fig5_two_mechanisms.png",6.5),
     7:("fig6_aliabadi_vel.png",4.9),8:("figY_elhawaz_ke.png",4.8),9:("fig7_pvloop.png",4.6),
     10:("fig8_flow_vs_work.png",6.5),11:("figZ_gc_vpeak.png",4.7),12:("fig9_evidence_hierarchy.png",6.2)}
import re
paras=d.paragraphs
replaced=[]
for i,p in enumerate(paras):
    if p.runs and p.runs[0].bold:
        m=re.match(r'Figure (\d+)\.',p.runs[0].text)
        if m:
            n=int(m.group(1)); img=paras[i-1]   # image paragraph precedes caption
            for r in list(img.runs): r._element.getparent().remove(r._element)
            fname,w=NUM[n]
            img.add_run().add_picture(f"{OUT}/{fname}",width=Inches(w))
            replaced.append(n)
d.save(FN)
print("Replaced figures:",sorted(replaced))
d2=Document(FN)
print("Inline images:",len(d2.inline_shapes),"| Tables:",len(d2.tables))
import zipfile; print("Valid docx:",zipfile.is_zipfile(FN))
